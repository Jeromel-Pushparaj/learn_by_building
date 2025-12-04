from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_syllabus_document():
    doc = Document()

    # Define the data for Semesters I to V (Anna Univ Reg 2021 B.E. CSE)
    # Format: Code, Title, L, T, P, Credits
    semesters = {
        "SEMESTER I": [
            ["THEORY", "", "", "", "", ""],
            ["IP3151", "Induction Programme", "0", "0", "0", "0"],
            ["HS3151", "Professional English - I", "3", "0", "0", "3"],
            ["MA3151", "Matrices and Calculus", "3", "1", "0", "4"],
            ["PH3151", "Engineering Physics", "3", "0", "0", "3"],
            ["CY3151", "Engineering Chemistry", "3", "0", "0", "3"],
            ["GE3151", "Problem Solving and Python Programming", "3", "0", "0", "3"],
            ["GE3152", "Heritage of Tamils", "1", "0", "0", "1"],
            [
                "GE3171",
                "Problem Solving and Python Programming Laboratory",
                "0",
                "0",
                "4",
                "2",
            ],
            ["BS3171", "Physics and Chemistry Laboratory", "0", "0", "4", "2"],
            ["GE3172", "English Laboratory", "0", "0", "2", "1"],
            ["HS3252", "Professional English - II", "2", "0", "0", "2"],
            ["MA3251", "Statistics and Numerical Methods", "3", "1", "0", "4"],
            ["PH3256", "Physics for Information Science", "3", "0", "0", "3"],
            [
                "BE3251",
                "Basic Electrical and Electronics Engineering",
                "3",
                "0",
                "0",
                "3",
            ],
            ["GE3251", "Engineering Graphics", "2", "0", "4", "4"],
            ["CS3251", "Programming in C", "3", "0", "0", "3"],
            ["GE3252", "Tamils and Technology", "1", "0", "0", "1"],
            ["PRACTICALS", "", "", "", "", ""],
            ["GE3271", "Engineering Practices Laboratory", "0", "0", "4", "2"],
            ["CS3271", "Programming in C Laboratory", "0", "0", "4", "2"],
            [
                "GE3272",
                "Communication Laboratory / Foreign Language",
                "0",
                "0",
                "4",
                "2",
            ],
            ["MA3354", "Discrete Mathematics", "3", "1", "0", "4"],
            [
                "CS3351",
                "Digital Principles and Computer Organization",
                "3",
                "0",
                "2",
                "4",
            ],
            ["CS3352", "Foundations of Data Science", "3", "0", "0", "3"],
            ["CS3301", "Data Structures", "3", "0", "0", "3"],
            ["CS3391", "Object Oriented Programming", "3", "0", "0", "3"],
            ["CS3311", "Data Structures Laboratory", "0", "0", "3", "1.5"],
            ["CS3381", "Object Oriented Programming Laboratory", "0", "0", "3", "1.5"],
            ["CS3361", "Data Science Laboratory", "0", "0", "4", "2"],
            ["GE3361", "Professional Development", "0", "0", "2", "1"],
            ["CS3452", "Theory of Computation", "3", "0", "0", "3"],
            [
                "CS3491",
                "Artificial Intelligence and Machine Learning",
                "3",
                "0",
                "2",
                "4",
            ],
            ["CS3492", "Database Management Systems", "3", "0", "0", "3"],
            ["CS3401", "Algorithms", "3", "0", "2", "4"],
            ["CS3451", "Introduction to Operating Systems", "3", "0", "0", "3"],
            ["GE3451", "Environmental Sciences and Sustainability", "2", "0", "0", "2"],
            ["CS3461", "Operating Systems Laboratory", "0", "0", "3", "1.5"],
            ["CS3481", "Database Management Systems Laboratory", "0", "0", "3", "1.5"],
            ["CS3591", "Computer Networks", "3", "0", "2", "4"],
            ["CS3501", "Compiler Design", "3", "0", "2", "4"],
            ["CB3491", "Cryptography and Cyber Security", "3", "0", "0", "3"],
            ["CS3551", "Distributed Computing", "3", "0", "0", "3"],
            ["CCS354", "Network Security", "2", "0", "2", "3"],
            ["CCS335", "Cloud Computing", "2", "0", "0", "3"],
            ["MX3081", "Elements of Literature", "3", "0", "0", "0"],
            ["CCS356", "Object Oriented Software Engineering", "3", "0", "2", "4"],
            ["CS3691", "Embedded Systems and IoT", "3", "0", "2", "4"],
            [
                "OCE351",
                "Environmental and Social Impact Assessment",
                "3",
                "0",
                "0",
                "3",
            ],
            ["CCS367", "Storage Technologies", "3", "0", "0", "3"],
            ["CCS358", "Principles of Programming Language", "3", "0", "0", "3"],
            ["CCS352", "Multimedia and Animation", "2", "0", "2", "3"],
            [
                "MX3085",
                "Well being with Traditional Practices - Yoga, Ayurveda and Siddha",
                "3",
                "0",
                "0",
                "0",
            ],
            ["GE3791", "Human Values and Ethics", "2", "0", "0", "2"],
            ["GE3751", "Principles of Management", "3", "0", "0", "3"],
            ["OHS351", "English for Competitive Examination", "3", "0", "0", "3"],
            ["OMG355", "Multivariate Data Analysis", "3", "0", "0", "3"],
            ["CS3711", "Summer internship", "0", "0", "0", "2"],
            ["CS3811", "Project Work/Internship", "0", "0", "300", "10"],
        ],
    }

    for semester_name, data in semesters.items():
        # Add Title
        doc.add_heading(semester_name, level=1)

        # Create Table
        table = doc.add_table(
            rows=len(data) + 2, cols=8
        )  # Adjusted for your Doc's columns
        table.style = "Table Grid"

        # --- Header Row 1 ---
        row0 = table.rows[0].cells
        row0[0].text = "Course Code"
        row0[1].text = "Course Titles"
        row0[2].text = "Teaching & Learning Scheme"
        row0[2].merge(row0[5])  # Merge L,T,P,SL cols
        row0[6].text = "Total no.of Hours"
        row0[7].text = "Total Credits"

        # --- Header Row 2 ---
        row1 = table.rows[1].cells
        row1[2].text = "L"
        row1[3].text = "T"
        row1[4].text = "P"
        row1[5].text = "SL"

        # Merge vertical cells for Code, Title, etc.
        row0[0].merge(row1[0])
        row0[1].merge(row1[1])
        row0[6].merge(row1[6])
        row0[7].merge(row1[7])

        # --- Fill Data ---
        for i, row_data in enumerate(data):
            cells = table.rows[i + 2].cells

            # Code and Title
            cells[0].text = row_data[0]
            cells[1].text = row_data[1]

            if row_data[0] in ["THEORY", "PRACTICALS", "TOTAL"]:
                # Bold and styling for category rows
                cells[0].paragraphs[0].runs[0].font.bold = True

            # L, T, P
            cells[2].text = row_data[2]
            cells[3].text = row_data[3]
            cells[4].text = row_data[4]

            # SL (Self Learning) - Assuming 0 as it's not in standard PDF
            cells[5].text = "0" if row_data[2].isdigit() else ""

            # Total Hours Calculation (approx 15 weeks per sem)
            # Formula: (L + T + P) * 15
            if row_data[2].isdigit():
                l = int(float(row_data[2]))
                t = int(float(row_data[3]))
                p = int(float(row_data[4]))
                total_hours = (l + t + p) * 15
                cells[6].text = str(total_hours)
            else:
                cells[6].text = ""

            # Credits
            cells[7].text = row_data[5]

            # Center align numbers
            for j in range(2, 8):
                cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save("B.E_CSE_Syllabus_I-VII.docx")
    print("Document created successfully!")


if __name__ == "__main__":
    create_syllabus_document()
